# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

out_dir = Path(r"E:\OpenAIglasses_for_Navigation-main\blind\docs")
out_dir.mkdir(parents=True, exist_ok=True)
out_path = out_dir / "参赛作品说明书_blind版.docx"

doc = Document()
normal = doc.styles['Normal']
normal.font.name = '宋体'
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
normal.font.size = Pt(12)

p = doc.add_paragraph('参赛作品说明书')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = p.runs[0]
r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.size = Pt(22)

cover_lines = [
    '作品名称：Blind 智能出行导航系统（blind）',
    '学　　校：　　　　　　　　　　　　　（待填写）',
    '学　　院：　　　　　　　　　　　　　（待填写）',
    '专业班别：　　　　　　　　　　　　　（待填写）',
    '学生姓名：　　　　　　　　　　　　　（待填写）',
    '指导老师：　　　　　　　　　　　　　（待填写）',
    '完成时间：　　　　　　　　　　　　　2026年05月',
]
for line in cover_lines:
    doc.add_paragraph(line)

doc.add_page_break()
doc.add_paragraph('目  录')
for item in [
    '第一章 简介','1.1 项目简介','1.2 核心功能','1.2.1 实时盲道识别与交通灯识别','1.2.2 导航播报与语音协同','1.2.3 地图与AI助手',
    '1.3 技术优势','1.4 用户价值','第二章 设计原理','2.1 系统总体架构','2.2 感知与推理链路','2.3 导航决策与语音播报机制','2.4 移动端与后端通信机制',
    '第三章 队员分工','3.1 项目分工建议','第四章 创新点','4.1 项目创新点','第五章 实用点','5.1 项目实用价值','第六章 总结','参考文献']:
    doc.add_paragraph(item)

def h1(t):
    p = doc.add_paragraph(t)
    p.style = doc.styles['Heading 1']

def h2(t):
    p = doc.add_paragraph(t)
    p.style = doc.styles['Heading 2']

def h3(t):
    p = doc.add_paragraph(t)
    p.style = doc.styles['Heading 3']

def body(t):
    doc.add_paragraph(t)

h1('第一章 简介')
h2('1.1 项目简介')
body('Blind 智能出行导航系统（blind）是一套面向视障与弱视人群的实时导航辅助方案，采用“手机端采集 + 后端视觉推理 + 实时语音反馈”的架构，帮助用户在城市道路和复杂环境中获得可执行、可理解的安全提示。')
body('项目覆盖 Android 端、FastAPI 后端、地图与 AI 助手能力，以及轻量 Web 控制台。系统通过统一鉴权机制支持移动端 Token 登录与 Web Session 登录，兼顾用户体验与接口安全。')

h2('1.2 核心功能')
h3('1.2.1 实时盲道识别与交通灯识别')
body('系统接收移动端相机实时帧，通过 YOLO 系列模型完成盲道分割与红绿灯识别，并持续输出“盲道状态 + 灯态状态 + 导航建议”。在导航模式下，系统优先保证安全语义，红灯相关播报具备更高优先级。')
h3('1.2.2 导航播报与语音协同')
body('后端根据规则引擎生成 guidance 文本，并通过 /ws/audio 推送 PCM 音频流；Android 端实现“导航优先、助手延后”的语音仲裁，避免导航播报与 AI 助手朗读重叠造成信息干扰。')
h3('1.2.3 地图与 AI 助手')
body('系统集成路线规划、周边检索、地理编码与逆地理编码接口，支持“定位—搜索—路线”闭环；AI 助手支持文本问答与语音输入转文本，提供导航、出行与日常问答的辅助能力。')

h2('1.3 技术优势')
body('（1）低耦合：Android、后端、模型推理、语音规则与鉴权模块相对独立，便于迭代。')
body('（2）高实时：采用 WebSocket 双向通道实现图像上传、引导文本下发和音频流播报。')
body('（3）可扩展：已预留 Web 端用户界面、ASR 代理、家属端扩展与多端协同空间。')

h2('1.4 用户价值')
body('项目将“感知、理解、播报”整合为一条可落地链路，使用户在不依赖复杂操作的前提下获得方向性与安全性提示；同时通过 AI 助手降低学习成本，提升系统可用性与长期使用意愿。')

h1('第二章 设计原理')
h2('2.1 系统总体架构')
body('系统由四层组成：设备采集层（Android Camera/Mic）、传输层（WebSocket + HTTP）、智能决策层（模型推理 + 导航编排 + 语音规则）、交互层（导航页、地图页、AI 助手页、Web 控制台）。')
h2('2.2 感知与推理链路')
body('Android 端将相机帧经压缩后发送至 /ws/camera；后端按帧执行盲道与交通灯推理，生成结构化结果。为平衡算力与时延，系统采用分帧发送、队列保护与降质策略，在拥塞时自动调节图像质量与发送节奏。')
h2('2.3 导航决策与语音播报机制')
body('NavigationOrchestrator 维护 IDLE、BLIND_NAV、TRAFFIC_TEST 三态。IDLE 状态不播报，避免误触发；BLIND_NAV 状态下综合盲道与交通灯结果生成建议文本，并通过节流与去重抑制重复播报。VoiceRuleEngine 将文本映射至本地音频资源，实现稳定、低延迟语音输出。')
h2('2.4 移动端与后端通信机制')
body('系统核心接口包括 /ws/camera、/ws/guidance、/ws/audio、/ws/viewer、/ws/asr_proxy 以及 /api/nav/start、/api/nav/stop、/api/assistant/chat、/api/map/*。鉴权采用“Web Session + 移动端 Token”并存，支持 app 登录态持久化与接口访问控制。')

h1('第三章 队员分工')
h2('3.1 项目分工建议')
h3('3.1.1 移动端与交互组')
body('负责 Android 页面结构、导航流程、语音输入、语音仲裁、异常提示与可视化交互优化；保障“可点、可读、可恢复”。')
h3('3.1.2 后端与算法组')
body('负责 FastAPI 服务、模型推理编排、WebSocket 稳定性、鉴权与日志、地图能力对接；持续优化时延、稳定性与错误恢复。')
h3('3.1.3 测试与产品组')
body('负责场景测试、回归测试、体验评估、文档维护与发布管理；建立测试清单（断网、弱网、重连、播报冲突、权限异常）。')

h1('第四章 创新点')
h2('4.1 项目创新点')
h3('4.1.1 安全优先的导航决策机制')
body('通过状态机与规则编排，将红绿灯语义、盲道语义和空闲态策略统一管理，降低“误报/错报”对用户行为的干扰风险。')
h3('4.1.2 导航与助手语音协同机制')
body('采用“导航优先”的语音仲裁方案，在同一终端内处理导航播报与助手朗读冲突，显著提升信息清晰度与可执行性。')
h3('4.1.3 轻量多端协同')
body('除 Android 主端外，系统内置轻量 Web 控制台用于状态观察与文本交互，降低部署和演示门槛，利于教学与比赛展示。')

h1('第五章 实用点')
h2('5.1 项目实用价值')
h3('5.1.1 即时性')
body('基于 WebSocket 的实时链路支持“边采集、边推理、边播报”，满足行走场景对时效性的要求。')
h3('5.1.2 可维护性')
body('模块化架构使模型、语音、地图、鉴权、前端页面可独立升级，便于后续接入家属端、多角色权限与更多传感器。')
h3('5.1.3 可落地性')
body('工程已具备端到端运行条件，包含依赖清单、启动脚本、接口契约与登录体系，具备持续迭代基础。')

h1('第六章 总结')
body('Blind 智能出行导航系统围绕“安全出行”核心目标，构建了从视觉感知到语音反馈的完整闭环。项目在实时通信、导航决策、语音协同与多端接入方面形成了工程化能力，既具备比赛展示价值，也具备继续产品化演进的技术基础。后续可进一步拓展家属协同、轨迹共享、后台告警与模型轻量化部署，持续提升真实场景可用性。')

h1('参考文献')
for i, r in enumerate([
    'Ultralytics YOLO Documentation. https://docs.ultralytics.com/',
    'FastAPI Documentation. https://fastapi.tiangolo.com/',
    'WebSocket Protocol (RFC 6455). https://datatracker.ietf.org/doc/html/rfc6455',
    'Android Developers Documentation. https://developer.android.com/',
    'Baidu Map Open Platform Documentation. https://lbsyun.baidu.com/'
], 1):
    body(f'[{i}] {r}')

doc.save(out_path)
print(str(out_path))
